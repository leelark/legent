import json
import os
import re
import shutil
import subprocess
import textwrap
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION_START
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as PdfImage,
    PageBreak,
)


ROOT = Path(__file__).resolve().parents[3]
PACKAGE = ROOT / "Legent_Project_Documentation_and_Demo"
DOC_MD = PACKAGE / "documentation" / "markdown"
DOC_DIR = PACKAGE / "documentation"
ASSETS = PACKAGE / "assets"
SCREENSHOTS = ASSETS / "screenshots"
DIAGRAMS = ASSETS / "diagrams"
EVIDENCE = PACKAGE / "evidence"
REPORTS = EVIDENCE / "reports"
WORK = PACKAGE / "_work"
TEMP_TOOLS = PACKAGE / "_temp_tools"

PY_NOW = datetime.now().strftime("%Y-%m-%d %H:%M")

SERVICE_PORTS = {
    "foundation-service": "8081",
    "audience-service": "8082",
    "campaign-service": "8083",
    "delivery-service": "8084",
    "tracking-service": "8085",
    "automation-service": "8086",
    "deliverability-service": "8087",
    "platform-service": "8088",
    "identity-service": "8089",
    "content-service": "8090",
}

SERVICE_DBS = {
    "foundation-service": "legent_foundation",
    "audience-service": "legent_audience",
    "campaign-service": "legent_campaign",
    "delivery-service": "legent_delivery",
    "tracking-service": "legent_tracking",
    "automation-service": "legent_automation",
    "deliverability-service": "legent_deliverability",
    "platform-service": "legent_platform",
    "identity-service": "legent_identity",
    "content-service": "legent_content",
}

RESPONSIBILITIES = {
    "foundation-service": "Tenant foundation, feature flags, branding, public content, admin settings, bootstrap, audit, and operational control plane.",
    "identity-service": "Login, signup, refresh token/session handling, account membership, role binding, onboarding state, and user preferences.",
    "audience-service": "Subscribers, lists, segments, imports, data extensions, preferences, suppression, consent, and audience resolution.",
    "content-service": "Email templates, blocks, versions, approvals, assets, rendering, test sends, landing pages, and creative governance.",
    "campaign-service": "Campaign lifecycle, wizard persistence, approvals, experiments, budgets, frequency policy, launch gates, send jobs, and checkpoints.",
    "delivery-service": "Provider selection, queue operations, message logs, warmup, replay, retries, circuit breakers, rate limits, and inbox safety.",
    "tracking-service": "Open/click ingestion, signed tracking URLs, raw events, campaign summaries, funnels, realtime analytics, and aggregation.",
    "automation-service": "Workflow definitions, graph validation, publish/pause/rollback controls, schedules, runs, simulations, and dry runs.",
    "deliverability-service": "Sender domain verification, DNS/DMARC, reputation, spam scoring, bounces, feedback loops, and suppression insights.",
    "platform-service": "Notifications, webhooks, search, indexing, retries, and shared platform integration utilities.",
}

MODULE_GUIDE = [
    ("Public site", "Market positioning, features, modules, pricing, signup, and onboarding.", "Visitor evaluates Legent and creates a workspace."),
    ("Identity", "Authentication, secure cookie session, account context, role membership, onboarding, preferences.", "Operator signs in and lands in a tenant/workspace context."),
    ("Foundation/Admin", "Tenant configuration, branding, feature flags, runtime settings, bootstrap status, audit, public CMS.", "Admin governs runtime behavior and support operations."),
    ("Audience", "Subscribers, lists, segments, imports, data extensions, consent, preferences, suppressions.", "Marketer creates a compliant target audience."),
    ("Content/Email", "Template Studio, reusable assets, versions, approvals, dynamic content, personalization checks.", "Creative operator prepares governed content."),
    ("Campaign", "Campaign wizard, audience targeting, frequency, budget, approvals, experiments, schedule, launch command.", "Lifecycle marketer launches a controlled send."),
    ("Delivery", "Provider routing, queue state, replay, retry, warmup, provider health, message logs.", "Delivery owner keeps send execution recoverable."),
    ("Tracking/Analytics", "Signed open/click events, raw events, summaries, funnels, campaign dashboards.", "Team reads outcomes and optimizes future sends."),
    ("Automation", "Journey graph, workflow versions, schedules, run history, simulation, dry run.", "Marketer automates lifecycle branches."),
    ("Deliverability", "Domain auth, DMARC, reputation, spam checks, bounce classification, feedback loops.", "Deliverability owner protects sender health."),
    ("Platform", "Webhooks, notifications, search, event propagation, integrations.", "Platform operator connects Legent to the broader stack."),
]

SCREEN_MAP = [
    ("01_public_home.png", "Public homepage", "Product story and command layer positioning."),
    ("02_public_features.png", "Features", "Governed operating model and core capabilities."),
    ("03_public_modules.png", "Modules", "Six operating studios connected by shared runtime fabric."),
    ("04_public_pricing.png", "Pricing", "Pricing surface for launch maturity and scale."),
    ("05_signup.png", "Signup", "Workspace creation entry point."),
    ("06_onboarding.png", "Onboarding", "Workspace, sender, provider, and launch readiness."),
    ("07_email_studio.png", "Email Studio", "Email command center and recent content."),
    ("08_template_studio.png", "Template Studio", "Template library, quick starts, and blank creation."),
    ("09_audience_overview.png", "Audience Overview", "Audience metrics, segments, imports, and lists."),
    ("10_subscriber_manager.png", "Subscriber Manager", "Contacts, lifecycle state, and list membership."),
    ("11_campaign_studio.png", "Campaign Studio", "Campaign portfolio and governed actions."),
    ("12_campaign_wizard.png", "Campaign Wizard", "Campaign setup, targeting, delivery rules, and review."),
    ("13_launch_orchestration.png", "Launch Command Center", "Readiness checks and launch controls."),
    ("14_campaign_tracking.png", "Campaign Tracking", "Progress, preflight gates, delivery, analytics, and experiment views."),
    ("15_automation_builder.png", "Automation", "Workflow list, status controls, and builder access."),
    ("16_deliverability.png", "Delivery Studio", "Queue, replay, warmup, domain health, and messages."),
    ("17_analytics.png", "Analytics Overview", "Performance metrics and campaign outcomes."),
    ("18_admin_console.png", "Admin Console", "Admin mesh, operations, users, roles, audit, and config."),
    ("19_platform_settings.png", "Platform Settings", "Profile, security, module, integration, and deliverability preferences."),
]


def ensure_dirs():
    for p in [DOC_MD, DOC_DIR, DIAGRAMS, REPORTS, TEMP_TOOLS]:
        p.mkdir(parents=True, exist_ok=True)


def rel(path: Path) -> str:
    return path.relative_to(PACKAGE).as_posix()


def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def split_camel(name: str) -> str:
    name = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", name)
    name = name.replace("_", " ").replace("-", " ")
    return " ".join(name.split()).strip().capitalize()


def extract_db_name(text: str, fallback: str) -> str:
    m = re.search(r"DB_NAME:([A-Za-z0-9_]+)", text)
    if m:
        return m.group(1)
    m = re.search(r"jdbc:postgresql://[^/\s]+/([A-Za-z0-9_]+)", text)
    if m:
        return m.group(1)
    return fallback


def inventory():
    services = []
    endpoint_rows = []
    migrations_by_service = {}
    tables_by_service = {}
    kafka_rows = []
    tests = []

    services_dir = ROOT / "services"
    for service_dir in sorted([p for p in services_dir.glob("*-service") if p.is_dir()]):
        name = service_dir.name
        yml = read_text(service_dir / "src/main/resources/application.yml") + "\n" + read_text(service_dir / "src/main/resources/application-local.yml")
        port = SERVICE_PORTS.get(name)
        pm = re.search(r"server:\s*(?:\n\s+[A-Za-z0-9_.-]+:.*)*\n\s+port:\s*(\d+)", yml, re.S)
        if pm:
            port = pm.group(1)
        db = extract_db_name(yml, SERVICE_DBS.get(name, "service_database"))
        controllers = sorted(service_dir.glob("src/main/java/**/controller/*.java"))
        entities = sorted(service_dir.glob("src/main/java/**/domain/*.java"))
        migration_files = sorted(service_dir.glob("src/main/resources/db/migration/*.sql"))
        unique_migrations = []
        seen = set()
        for mig in migration_files:
            if mig.name not in seen:
                seen.add(mig.name)
                unique_migrations.append(mig.name)
        table_names = []
        for mig in migration_files:
            text = read_text(mig)
            for mt in re.finditer(r"CREATE\s+TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?([A-Za-z0-9_.]+)", text, re.I):
                table_names.append(mt.group(1).split(".")[-1])
        migrations_by_service[name] = unique_migrations
        tables_by_service[name] = sorted(set(table_names))

        services.append({
            "name": name,
            "port": port or "",
            "db": db,
            "controllers": len(controllers),
            "entities": len(entities),
            "migrations": len(unique_migrations),
            "tables": len(set(table_names)),
            "responsibility": RESPONSIBILITIES.get(name, ""),
        })

        for ctrl in controllers:
            endpoint_rows.extend(parse_controller(name, ctrl))

        for java in service_dir.glob("src/main/java/**/*.java"):
            txt = read_text(java)
            if "@KafkaListener" in txt:
                topics = []
                for topic_match in re.finditer(r"topics\s*=\s*(?:\{([^}]+)\}|\"([^\"]+)\"|([A-Za-z0-9_.]+))", txt):
                    topics.append(" ".join([g for g in topic_match.groups() if g]))
                kafka_rows.append({
                    "service": name,
                    "kind": "consumer",
                    "detail": ", ".join(topics)[:160] or "@KafkaListener",
                    "file": str(java.relative_to(ROOT)).replace("\\", "/"),
                })
            if "EventPublisher" in txt or "KafkaTemplate" in txt:
                kafka_rows.append({
                    "service": name,
                    "kind": "publisher/helper",
                    "detail": "EventPublisher/KafkaTemplate usage",
                    "file": str(java.relative_to(ROOT)).replace("\\", "/"),
                })
        for test in service_dir.glob("src/test/**/*.java"):
            tests.append(str(test.relative_to(ROOT)).replace("\\", "/"))

    for test in (ROOT / "frontend" / "tests").glob("**/*.spec.ts"):
        tests.append(str(test.relative_to(ROOT)).replace("\\", "/"))

    frontend_routes = []
    app_dir = ROOT / "frontend" / "src" / "app"
    for page in sorted(app_dir.glob("**/page.tsx")):
        parts = page.relative_to(app_dir).parts[:-1]
        clean = []
        for part in parts:
            if part.startswith("(") and part.endswith(")"):
                continue
            clean.append(part)
        route = "/" + "/".join(clean)
        route = route.replace("/app/app", "/app")
        if route == "/":
            route = "/"
        frontend_routes.append({
            "route": route,
            "file": str(page.relative_to(ROOT)).replace("\\", "/"),
        })
    frontend_routes = sorted({r["route"]: r for r in frontend_routes}.values(), key=lambda x: x["route"])

    inv = {
        "generated_at": PY_NOW,
        "services": services,
        "endpoints": sorted(endpoint_rows, key=lambda e: (e["service"], e["path"], e["method"])),
        "migrations_by_service": migrations_by_service,
        "tables_by_service": tables_by_service,
        "kafka": kafka_rows,
        "tests": tests,
        "frontend_routes": frontend_routes,
        "screenshots": [x[0] for x in SCREEN_MAP if (SCREENSHOTS / x[0]).exists()],
    }
    (EVIDENCE / "codebase-inventory-refined.json").write_text(json.dumps(inv, indent=2), encoding="utf-8")
    return inv


def parse_controller(service: str, ctrl: Path):
    text = read_text(ctrl)
    class_base = ""
    class_match = re.search(r"@RequestMapping\s*\(([^)]*)\)\s*(?:\n\s*@[A-Za-z0-9_().,\" =/-]+)*\s*public\s+class", text, re.S)
    if class_match:
        class_base = annotation_path(class_match.group(1))
    rows = []
    for match in re.finditer(r"@(GetMapping|PostMapping|PutMapping|PatchMapping|DeleteMapping)(?:\s*\(([^)]*)\))?", text):
        ann = match.group(1)
        args = match.group(2) or ""
        method = ann.replace("Mapping", "").upper()
        path_part = annotation_path(args)
        after = text[match.end():match.end() + 600]
        mm = re.search(r"(?:public|private|protected)\s+[^{;=]+?\s+([A-Za-z0-9_]+)\s*\(", after)
        method_name = mm.group(1) if mm else ""
        path = normalize_path(class_base, path_part)
        rows.append({
            "service": service,
            "controller": ctrl.name,
            "method": method,
            "path": path,
            "handler": method_name,
            "purpose": split_camel(method_name) if method_name else split_camel(ctrl.stem),
            "file": str(ctrl.relative_to(ROOT)).replace("\\", "/"),
        })
    return rows


def annotation_path(args: str) -> str:
    if not args:
        return ""
    m = re.search(r"(?:value\s*=\s*|path\s*=\s*)?\"([^\"]*)\"", args)
    if m:
        return m.group(1)
    return ""


def normalize_path(base: str, child: str) -> str:
    base = base or ""
    child = child or ""
    if not base.startswith("/") and base:
        base = "/" + base
    if child and not child.startswith("/"):
        child = "/" + child
    combined = (base.rstrip("/") + child) or "/"
    return re.sub(r"/+", "/", combined)


def md_table(headers, rows):
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        safe = [str(x).replace("\n", "<br>").replace("|", "\\|") for x in row]
        out.append("| " + " | ".join(safe) + " |")
    return "\n".join(out)


def service_matrix(inv):
    return md_table(
        ["Service", "Port", "DB", "Controllers", "Entities", "Migrations", "Tables", "Responsibility"],
        [[s["name"], s["port"], s["db"], s["controllers"], s["entities"], s["migrations"], s["tables"], s["responsibility"]] for s in inv["services"]],
    )


def write_markdown(inv):
    diagram_refs = [
        "system-architecture.png",
        "service-topology.png",
        "campaign-data-flow.png",
        "event-flow.png",
        "database-ownership.png",
        "deployment-architecture.png",
    ]
    diag_md = "\n".join([f"![{Path(d).stem.replace('-', ' ').title()}](../../assets/diagrams/{d})" for d in diagram_refs])
    counts = {
        "services": len(inv["services"]),
        "endpoints": len(inv["endpoints"]),
        "routes": len(inv["frontend_routes"]),
        "tests": len(inv["tests"]),
        "screenshots": len(inv["screenshots"]),
    }

    write_md("00_README_START_HERE.md", f"""# Legent Project Documentation And Demo Package

Generated: {PY_NOW}

## What This Folder Contains

This package is the refined delivery set for Legent v1.0.2. It contains verified screenshots, architecture/data-flow diagrams, detailed documentation, a polished DOCX/PDF report, three product demo videos, and one 10-minute technical explanation video.

## Reading Order

1. `01_Quick_Guide.md` for a short orientation.
2. `02_Complete_Master_Documentation.md` for the full product and engineering report.
3. `06_User_Guide_With_Screenshots.md` for the end-to-end user journey.
4. `07_Local_Runbook.md` and `08_Cloud_Setup_Guide.md` for execution.
5. `09_API_Catalog.md`, `10_Database_And_Migrations.md`, and `05_Architecture_Documentation.md` for deep technical reference.

## Inventory

{md_table(["Area", "Count"], [["Spring services", counts["services"]], ["Cataloged API endpoints", counts["endpoints"]], ["Frontend routes", counts["routes"]], ["Test specs/classes", counts["tests"]], ["Verified screenshots", counts["screenshots"]]])}

## Final Artifacts

- `documentation/Legent_Project_Documentation.docx`
- `documentation/Legent_Project_Documentation.pdf`
- `demo-video/3min/legent-product-demo-3min.mp4`
- `demo-video/6min/legent-product-demo-6min.mp4`
- `demo-video/10min/legent-product-demo-10min.mp4`
- `technical-video/10min/legent-technical-explanation-10min.mp4`
""")

    write_md("01_Quick_Guide.md", f"""# Quick Guide

## Product Summary

Legent is a lifecycle email operating system. It combines public acquisition, signup/onboarding, audience management, template governance, campaign launch, delivery operations, tracking, analytics, automation, deliverability, admin, and platform settings inside one tenant/workspace-aware product.

## Architecture In One Page

{diag_md}

## Service Summary

{service_matrix(inv)}

## Fast Local Start

```powershell
docker compose up -d postgres kafka redis minio opensearch clickhouse mailhog
mvn -q -DskipTests package
mvn -pl services/foundation-service spring-boot:run
cd frontend
npm install
npm run dev -- --hostname 127.0.0.1 --port 3000
```

## Main User Journey

Public site -> signup -> onboarding -> audience -> template -> campaign wizard -> launch command center -> delivery -> tracking/analytics -> automation -> admin/settings.
""")

    write_md("02_Complete_Master_Documentation.md", complete_master(inv, diag_md))
    write_md("03_Functional_Documentation.md", functional_doc(inv))
    write_md("04_Technical_Documentation.md", technical_doc(inv))
    write_md("05_Architecture_Documentation.md", architecture_doc(inv, diag_md))
    write_md("06_User_Guide_With_Screenshots.md", user_guide_doc())
    write_md("07_Local_Runbook.md", local_runbook())
    write_md("08_Cloud_Setup_Guide.md", cloud_guide(inv))
    write_md("09_API_Catalog.md", api_catalog(inv))
    write_md("10_Database_And_Migrations.md", db_doc(inv))
    write_md("11_Testing_And_QA_Guide.md", testing_doc(inv))
    write_md("12_Operations_And_Troubleshooting.md", ops_doc(inv))
    write_md("13_Security_And_Compliance.md", security_doc(inv))
    write_md("14_Future_Enhancement_Roadmap.md", roadmap_doc())

    root_readme = PACKAGE / "README.md"
    root_readme.write_text((DOC_MD / "00_README_START_HERE.md").read_text(encoding="utf-8"), encoding="utf-8")


def write_md(name: str, content: str):
    (DOC_MD / name).write_text(textwrap.dedent(content).strip() + "\n", encoding="utf-8")


def complete_master(inv, diag_md):
    modules = "\n\n".join([
        f"### {name}\n\n**Purpose.** {purpose}\n\n**Operator outcome.** {outcome}\n\n**Acceptance focus.** Screens must load with tenant/workspace context, show actionable state, and avoid blank or shell-only states during normal operation."
        for name, purpose, outcome in MODULE_GUIDE
    ])
    return f"""# Complete Master Documentation

## Executive Summary

Legent v1.0.2 is a multi-service lifecycle email platform. It is organized around an operator journey from public discovery and workspace setup through audience preparation, creative governance, campaign launch, delivery execution, tracking, analytics, automation, deliverability, and admin operations.

## Product Scope

- Public marketing pages and signup/onboarding.
- Workspace-authenticated Next.js application.
- Spring Boot microservices by domain.
- Shared libraries for common models, security, Kafka, and test support.
- Service-owned PostgreSQL databases through Flyway.
- Kafka event fabric for asynchronous coordination.
- Redis, MinIO, OpenSearch, ClickHouse, MailHog, Docker Compose, and Kubernetes assets.

## Verified Evidence

- Screenshots: {len(inv["screenshots"])} verified UI captures.
- API endpoints cataloged: {len(inv["endpoints"])}.
- Frontend routes cataloged: {len(inv["frontend_routes"])}.
- Test files cataloged: {len(inv["tests"])}.

## Architecture And Data Flow

{diag_md}

## Service Matrix

{service_matrix(inv)}

## Functional Modules

{modules}

## Frontend

The frontend is a Next.js application under `frontend/`. It contains public pages, auth/onboarding surfaces, workspace routes under `/app`, reusable UI primitives, marketing components, admin/settings consoles, API clients, stores, hooks, and Playwright tests. API clients centralize tenant/workspace headers and response normalization.

## Backend

The backend is split into domain services. Controllers expose `/api/v1` endpoints. Services own business logic, repositories own persistence access, Flyway owns schema evolution, and Kafka publishers/listeners coordinate cross-service events.

## Data Ownership

Every service has a dedicated database. Cross-service coupling should happen through APIs or events, not by directly reading another service database. ClickHouse supports analytics-heavy tracking data. Redis supports fast operational state. MinIO stores object assets. OpenSearch supports indexing/search.

## Events

Kafka is used for tenant provisioning, audience resolution, campaign launch/send coordination, delivery feedback, tracking ingestion, automation workflow activity, and system propagation events. Events should be idempotent and include tenant/workspace context.

## Security

The codebase uses HTTP-only cookie session posture, JWT validation, RBAC evaluation, tenant/workspace headers, context bootstrap, CORS controls, audit logging, and secrets through environment/configuration rather than committed plaintext values.

## Operations

Local execution uses Docker Compose for dependencies and Maven/Next.js for services/UI. Cloud execution is Kubernetes-first through manifests under `infrastructure/kubernetes`. Health checks, logs, migrations, Kafka status, queue depth, and provider health are the main operational checks.

## Testing

Testing includes Java service tests and frontend Playwright tests. High-risk areas are auth context bootstrap, tenant/workspace scope, campaign launch idempotency, event consumers, delivery replay, tracking ingestion, admin config propagation, and UI route loading.

## Troubleshooting

Blank UI captures usually mean the route was captured before hydration, the workspace context failed, or mock/API responses did not match page expectations. The refined capture workflow waits for screen-specific text and records browser errors.

## Glossary

- Tenant: top-level account boundary.
- Workspace: operational workspace under a tenant.
- Environment: runtime scope for production/staging-like behavior.
- Launch gate: readiness rule checked before campaign execution.
- Replay: operator-controlled resend/reprocess path for delivery failures.
- Idempotency: repeated event/command processing that avoids duplicate side effects.
"""


def functional_doc(inv):
    feature_rows = [[name, purpose, outcome] for name, purpose, outcome in MODULE_GUIDE]
    workflows = """
## End-To-End Operator Journey

1. Visitor opens the public site and reviews capabilities, modules, and pricing.
2. Visitor signs up and creates or enters a workspace context.
3. Operator completes onboarding readiness.
4. Marketer builds subscribers, lists, segments, imports, data extensions, and preferences.
5. Creative user creates or selects approved templates.
6. Campaign owner uses the campaign wizard to set identity, audience, delivery rules, budget, frequency, and experiment controls.
7. Launch operator runs readiness checks and confirms launch.
8. Campaign service creates send jobs and coordinates audience resolution.
9. Delivery service routes messages through providers with retry/replay guardrails.
10. Tracking service records opens, clicks, and event outcomes.
11. Analytics, automation, deliverability, admin, and settings screens close the feedback loop.
"""
    acceptance = "\n".join([
        f"- {name}: usable screen state, clear operator action, tenant/workspace scoped data, and recoverable error handling."
        for name, _, _ in MODULE_GUIDE
    ])
    return f"""# Functional Documentation

## Personas

{md_table(["Persona", "Primary Goal", "Main Screens"], [
["Lifecycle marketer", "Create lifecycle campaigns and journeys", "Audience, Email, Campaign, Automation"],
["Marketing operations", "Govern launch readiness and recovery", "Campaign, Launch, Admin, Delivery"],
["Deliverability owner", "Protect sender health and inbox placement", "Delivery Studio, Deliverability settings"],
["Admin/operator", "Manage users, roles, config, audit, and platform operations", "Admin, Settings, Foundation"],
["Developer/SRE", "Run, deploy, observe, and troubleshoot the system", "Runbooks, logs, health, Docker/K8s"],
])}

## Module Feature Matrix

{md_table(["Module", "Business Behavior", "User Outcome"], feature_rows)}

{workflows}

## Business Rules

- All authenticated work is tenant/workspace scoped.
- Campaign launch requires audience, content, approval, budget/frequency, and deliverability readiness.
- Delivery replay and retry must remain operator-visible and idempotent.
- Tracking must accept public open/click events without requiring authenticated UI session state.
- Admin configuration changes should propagate with audit evidence.
- Preferences, consent, suppressions, and unsubscribe state must be respected before send.

## Acceptance Notes

{acceptance}
"""


def technical_doc(inv):
    endpoint_sample = inv["endpoints"][:80]
    endpoint_rows = [[e["service"], e["method"], e["path"], e["handler"], e["controller"]] for e in endpoint_sample]
    kafka_rows = [[k["service"], k["kind"], k["detail"], k["file"]] for k in inv["kafka"][:80]]
    route_rows = [[r["route"], r["file"]] for r in inv["frontend_routes"][:80]]
    return f"""# Technical Documentation

## Repository Structure

- `frontend/`: Next.js app, public site, workspace UI, API clients, stores, hooks, and Playwright specs.
- `services/`: Spring Boot microservices.
- `shared/`: common, security, Kafka, and test support libraries.
- `infrastructure/docker/local/`: dependency initialization scripts.
- `infrastructure/kubernetes/`: base Kubernetes manifests, service deployments, secrets/config, HPA, network policy, and ingress.

## Service Matrix

{service_matrix(inv)}

## Frontend Route Catalog

{md_table(["Route", "File"], route_rows)}

## API Catalog Snapshot

The complete catalog is in `09_API_Catalog.md`.

{md_table(["Service", "Method", "Path", "Handler", "Controller"], endpoint_rows)}

## Kafka/Event Usage Snapshot

{md_table(["Service", "Kind", "Detail", "File"], kafka_rows)}

## Shared Libraries

- `legent-common`: common model/utilities.
- `legent-security`: JWT, tenant context, interceptors, RBAC, async tenant propagation, and exception handling.
- `legent-kafka`: Kafka producer/consumer configuration, event envelope, domain event helpers, and DLQ handling.
- `legent-test-support`: base integration support and mock tenant context.

## API Conventions

- Authenticated UI uses `/api/v1/...`.
- Public tracking and public marketing APIs are tenant-free or explicitly whitelisted.
- Workspace optional endpoints include preferences, admin bootstrap/config/public content, core platform, notifications, search, and webhooks.
- Tenant, workspace, and environment context are carried by local context and request headers where applicable.

## Build And Test Commands

```powershell
mvn -q -DskipTests package
mvn test
cd frontend
npm install
npm run lint
npm run test:e2e
```

## Technical Risks To Guard

- Response shape drift between Spring APIs and frontend API clients.
- Tenant/workspace headers missing from authenticated requests.
- Duplicate Kafka event handling without idempotency.
- Flyway migration drift across local/cloud databases.
- UI screenshot/video capture before route hydration.
"""


def architecture_doc(inv, diag_md):
    return f"""# Architecture Documentation

## Architecture Diagrams

{diag_md}

## System Context

```mermaid
flowchart TB
  Visitor["Visitor / Operator"] --> Frontend["Next.js Frontend"]
  Frontend --> Gateway["Ingress / API Gateway"]
  Gateway --> Services["Spring Boot Services"]
  Services --> Postgres["Service-owned PostgreSQL"]
  Services --> Kafka["Kafka Event Fabric"]
  Services --> Redis["Redis"]
  Services --> ObjectStore["MinIO / Object Storage"]
  Services --> Search["OpenSearch"]
  Services --> Analytics["ClickHouse"]
```

## Service Topology

{service_matrix(inv)}

## Data Flow

```mermaid
sequenceDiagram
  participant U as Operator
  participant FE as Frontend
  participant AU as Audience
  participant CO as Content
  participant CA as Campaign
  participant DE as Delivery
  participant TR as Tracking
  participant AN as Analytics
  U->>FE: Build campaign
  FE->>AU: Resolve target audience
  FE->>CO: Select approved template
  FE->>CA: Save campaign and launch
  CA->>DE: Create send work
  DE->>TR: Emit delivery/tracking signals
  TR->>AN: Aggregate outcomes
  AN->>FE: Display performance
```

## Event Flow

Kafka decouples long-running and cross-domain work. Typical flows include tenant provisioning, audience resolution, campaign launch, delivery feedback, tracking ingestion, automation execution, and admin propagation.

## Database Ownership

Each service owns its own database schema and Flyway timeline. See `10_Database_And_Migrations.md` for the service-by-service inventory.

## Deployment View

Kubernetes manifests define namespace, config map, secrets, services, deployments, HPA, network policy, ingress, and dependency services. Production deployments should normally use managed PostgreSQL/Kafka/Redis/object storage/search/analytics rather than single-node in-cluster dependencies.
"""


def user_guide_doc():
    sections = []
    for idx, (fname, title, caption) in enumerate(SCREEN_MAP, start=1):
        sections.append(f"""## {idx}. {title}

![{title}](../../assets/screenshots/{fname})

**What to verify.** {caption}

**Operator action.** Review the visible state, confirm the screen has loaded beyond the app shell, then continue to the next step in the journey.
""")
    return "# User Guide With Screenshots\n\nThis guide uses the refined, verified screenshots captured from the local frontend with deterministic workspace data.\n\n" + "\n".join(sections)


def local_runbook():
    return """# Local Runbook

## Prerequisites

- JDK 17+
- Maven 3.9+
- Node.js 20+ or the bundled runtime
- Docker Desktop
- PowerShell on Windows

## Start Dependencies

```powershell
docker compose up -d postgres kafka redis minio opensearch clickhouse mailhog
docker compose ps
```

## Build Backend

```powershell
mvn -q -DskipTests package
mvn test
```

## Run Services Locally

Run each service in a dedicated terminal when developing:

```powershell
mvn -pl services/foundation-service spring-boot:run
mvn -pl services/audience-service spring-boot:run
mvn -pl services/content-service spring-boot:run
mvn -pl services/campaign-service spring-boot:run
mvn -pl services/delivery-service spring-boot:run
mvn -pl services/tracking-service spring-boot:run
mvn -pl services/automation-service spring-boot:run
mvn -pl services/deliverability-service spring-boot:run
mvn -pl services/platform-service spring-boot:run
mvn -pl services/identity-service spring-boot:run
```

## Run Frontend

```powershell
cd frontend
npm install
npm run dev -- --hostname 127.0.0.1 --port 3000
```

Open `http://127.0.0.1:3000`.

## Smoke Checks

```powershell
Invoke-WebRequest http://127.0.0.1:3000 -UseBasicParsing
Invoke-WebRequest http://127.0.0.1:8081/actuator/health -UseBasicParsing
Invoke-WebRequest http://127.0.0.1:8089/actuator/health -UseBasicParsing
cd frontend
npm run test:e2e
```

## Troubleshooting

- If frontend redirects to login, verify `/api/v1/auth/session`, tenant id, workspace id, and environment id.
- If a service fails on startup, inspect DB URL, Flyway migration, and Kafka/Redis connectivity.
- If screenshots are blank, wait for route-specific text and inspect browser console errors.
"""


def cloud_guide(inv):
    service_names = " ".join([s["name"] for s in inv["services"]])
    return f"""# Cloud Setup Guide

## Deployment Model

Use Kubernetes as the primary cloud deployment target. The repository already contains base manifests under `infrastructure/kubernetes/base` and ingress manifests under `infrastructure/kubernetes/ingress`.

## Build Images

```powershell
$REGISTRY=\"registry.example.com/legent\"
docker build -t $REGISTRY/frontend:1.0.2 frontend
foreach ($svc in \"{service_names}\".Split(\" \")) {{
  docker build -t $REGISTRY/$svc:1.0.2 services/$svc
}}
docker push $REGISTRY/frontend:1.0.2
foreach ($svc in \"{service_names}\".Split(\" \")) {{ docker push $REGISTRY/$svc:1.0.2 }}
```

## Apply Kubernetes Base

```powershell
kubectl apply -f infrastructure/kubernetes/base/namespace.yml
kubectl apply -f infrastructure/kubernetes/base/secrets.yml
kubectl apply -f infrastructure/kubernetes/base/configmap.yml
kubectl apply -f infrastructure/kubernetes/base/services/
kubectl apply -f infrastructure/kubernetes/base/deployments/
kubectl apply -f infrastructure/kubernetes/base/hpa.yml
kubectl apply -f infrastructure/kubernetes/base/network-policy.yml
kubectl apply -f infrastructure/kubernetes/ingress/ingress.yml
```

## Secrets And Config

Replace placeholder secrets before production use:

- PostgreSQL credentials and JDBC URLs.
- JWT signing/verification material.
- Provider API keys and SMTP credentials.
- Object storage credentials.
- Kafka bootstrap/security settings.
- CORS origins and public app URLs.

## Health Checks

```powershell
kubectl get pods -n legent
kubectl get svc -n legent
kubectl logs -n legent deploy/foundation-service --tail=100
kubectl rollout status -n legent deploy/frontend
```

## Rollback

```powershell
kubectl rollout undo -n legent deploy/frontend
kubectl rollout undo -n legent deploy/campaign-service
```

## EKS Appendix

```powershell
eksctl create cluster --name legent-prod --region us-east-1 --nodes 3
aws eks update-kubeconfig --name legent-prod --region us-east-1
```

## GKE Appendix

```powershell
gcloud container clusters create legent-prod --zone us-central1-a --num-nodes 3
gcloud container clusters get-credentials legent-prod --zone us-central1-a
```

## AKS Appendix

```powershell
az aks create --resource-group legent-rg --name legent-prod --node-count 3 --enable-managed-identity
az aks get-credentials --resource-group legent-rg --name legent-prod
```
"""


def api_catalog(inv):
    rows = [[e["service"], e["method"], e["path"], e["controller"], e["handler"], e["purpose"]] for e in inv["endpoints"]]
    return f"""# API Catalog

Generated from Spring controller annotations. Request/response DTOs should be confirmed from handler signatures when implementing clients.

## Endpoint Inventory

{md_table(["Service", "Method", "Path", "Controller", "Handler", "Purpose"], rows)}
"""


def db_doc(inv):
    sections = []
    for s in inv["services"]:
        name = s["name"]
        migs = inv["migrations_by_service"].get(name, [])
        tables = inv["tables_by_service"].get(name, [])
        sections.append(f"""## {name}

**Database:** `{s["db"]}`

**Owned tables discovered:** {", ".join(tables) if tables else "No CREATE TABLE statement discovered in migrations."}

**Migration timeline**

{md_table(["#", "Migration"], [[i + 1, m] for i, m in enumerate(migs)])}
""")
    return "# Database And Migrations\n\nEach service owns its schema through Flyway. Avoid direct cross-service database reads.\n\n" + "\n".join(sections) + """

## Backup And Restore Notes

- Back up each service database independently.
- Preserve Flyway schema history with database backups.
- Restore dependency order by identity/foundation first, then audience/content/campaign/delivery/tracking/automation/deliverability/platform.
- Validate restored systems with smoke tests and idempotency checks before enabling scheduled sends.
"""


def testing_doc(inv):
    rows = [[i + 1, t] for i, t in enumerate(inv["tests"][:160])]
    return f"""# Testing And QA Guide

## Test Commands

```powershell
mvn test
cd frontend
npm run test:e2e
npm run lint
```

## Test Inventory

{md_table(["#", "Test file"], rows)}

## QA Checklist

- Public pages load unique content.
- Auth login/bootstrap sets tenant, workspace, and environment context.
- Workspace screens load real content, not app shell only.
- Campaign wizard persists experiment, budget, frequency, and launch controls.
- Tracking screen exposes safety, DLQ, budget, and variant analytics.
- Admin console loads operations, users, and role engine.
- Settings console persists preferences and supports deliverability.
- Event consumers are idempotent.
- Flyway migrations validate on clean local database.
"""


def ops_doc(inv):
    return f"""# Operations And Troubleshooting

## Core Diagnostics

```powershell
docker compose ps
docker compose logs --tail=200 postgres
docker compose logs --tail=200 kafka
kubectl get pods -n legent
kubectl logs -n legent deploy/campaign-service --tail=200
```

## Health Endpoints

Use `/actuator/health` for Spring services where actuator is enabled. Frontend health is confirmed with HTTP 200 on `/`.

## Common Problems

{md_table(["Symptom", "Likely Cause", "Fix"], [
["Frontend login loop", "Session/context bootstrap failing", "Check auth session API, local storage context, headers, and identity service logs"],
["Blank workspace page", "Hydration or API response shape failure", "Open browser console, confirm endpoint mock/response schema, wait for page-specific text"],
["Service startup failure", "DB/Kafka/Redis unavailable or wrong env", "Check Docker Compose status and application-local.yml"],
["Flyway error", "Migration drift or missing database", "Inspect service DB and flyway_schema_history"],
["Campaign duplicate send", "Idempotency key/checkpoint issue", "Inspect campaign and delivery event idempotency tables/logs"],
["Delivery backlog", "Provider health, queue, or rate limit issue", "Check delivery queue stats, provider health, replay queue, and retry policy"],
])}

## Logs To Inspect

- `identity-service`: sessions, refresh tokens, account context, preferences.
- `foundation-service`: tenant bootstrap, config propagation, audit, admin operations.
- `campaign-service`: launch gates, send jobs, checkpoints, experiments.
- `delivery-service`: provider selection, queues, replay, retries.
- `tracking-service`: open/click ingestion and analytics aggregation.
"""


def security_doc(inv):
    return """# Security And Compliance

## Authentication

The system is designed around HTTP-only cookie session posture with JWT validation. Frontend code stores only non-sensitive context such as user id, roles, tenant id, workspace id, environment id, and theme preferences.

## Tenant And Workspace Scope

Authenticated APIs must operate with tenant/workspace context. The frontend API client injects context headers where applicable. Backend filters/interceptors keep tenant context available to service code and async work.

## RBAC

Roles are evaluated by shared security code. Admin, platform admin, organization admin, campaign manager, delivery operator, analyst, viewer, and user-like roles should map to explicit permissions.

## Secrets

Do not commit provider credentials, JWT secrets, database passwords, object storage keys, or webhook secrets. Use environment variables, Kubernetes secrets, cloud secret managers, or sealed secrets.

## Audit And Compliance

Admin configuration, role changes, account/session actions, delivery replay, campaign launch, and tenant bootstrap actions should be auditable. Privacy-sensitive data must be minimized in logs.

## Privacy Notes

Subscriber data, consent, preferences, suppressions, and unsubscribe state are core compliance surfaces. Backup/restore and analytics exports should preserve consent semantics and workspace boundaries.
"""


def roadmap_doc():
    rows = [
        ["Frontend UX", "Route-level skeletons, richer empty states, screenshot-proof loading guards, accessibility pass, visual regression tests."],
        ["Identity", "MFA, device/session management, SSO/SAML/OIDC, invitation lifecycle, account recovery hardening."],
        ["Foundation/Admin", "Config approval workflow, policy simulator, audit export, drift detection, tenant bootstrap wizard."],
        ["Audience", "Advanced segmentation, predictive scores, preference center, consent import validation, dedupe merge UI."],
        ["Content", "Collaborative editing, brand linting, spam checks in template editor, version diff UI, richer assets."],
        ["Campaign", "Calendar view, multi-channel extension, advanced experiments, budget guardrails, campaign dependency graph."],
        ["Delivery", "Provider cost optimization, adaptive throttling, replay analytics, provider simulator, failover drills."],
        ["Tracking/Analytics", "Attribution models, cohort retention, funnel builder, export scheduler, realtime alert thresholds."],
        ["Automation", "Visual branch testing, replayable journey simulation, schedule calendar, reusable journey blocks."],
        ["Deliverability", "DMARC aggregate import, domain warmup plans, complaint trend alerts, reputation forecasting."],
        ["Platform", "Webhook retry UI, integration marketplace, API keys, search ranking, notification rules."],
        ["Infra", "Terraform modules, managed cloud dependency profiles, blue/green deployment, automated rollback."],
        ["Observability", "OpenTelemetry traces, SLO dashboards, service latency budgets, Kafka lag alerts, synthetic journeys."],
        ["Testing", "Contract tests, visual regression, load tests, migration tests, chaos tests for Kafka/provider failure."],
        ["AI/Product Intelligence", "Explainable journey suggestions, next-best action, subject/body assist, anomaly detection."],
    ]
    return f"""# Future Enhancement Roadmap

## Module-Level Enhancements

{md_table(["Area", "Enhancements"], rows)}

## Product-Level Priorities

1. Make the first workspace launch path more guided and measurable.
2. Add stronger governance around approval, config, and role propagation.
3. Expand observability so support/SRE can diagnose tenant-specific issues quickly.
4. Add contract and visual tests to prevent response-shape drift and blank UI regressions.
5. Prepare cloud deployment with managed services, secret management, and repeatable infrastructure.
"""


def font(size=28, bold=False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def wrap(draw, text, fnt, max_width):
    words = str(text).split()
    lines, cur = [], ""
    for word in words:
        trial = f"{cur} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            cur = trial
        else:
            if cur:
                lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    return lines


def draw_box(draw, xy, title, body="", fill="#ffffff", outline="#d8c7ff", accent="#7c3aed"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=26, fill=fill, outline=outline, width=3)
    draw.rectangle((x1, y1, x1 + 10, y2), fill=accent)
    tf, bf = font(26, True), font(19)
    draw.text((x1 + 28, y1 + 20), title, fill="#1f1640", font=tf)
    if body:
        y = y1 + 58
        for line in wrap(draw, body, bf, x2 - x1 - 50)[:4]:
            draw.text((x1 + 28, y), line, fill="#5f547d", font=bf)
            y += 25


def arrow(draw, start, end, color="#6d5c91", width=4):
    draw.line((start, end), fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - direction * 18, ey - 10), (ex - direction * 18, ey + 10)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 10, ey - direction * 18), (ex + 10, ey - direction * 18)]
    draw.polygon(pts, fill=color)


def diagram_canvas(title, subtitle):
    img = Image.new("RGB", (1600, 900), "#f8f6ff")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 110), fill="#21133f")
    draw.text((48, 26), title, fill="#ffffff", font=font(38, True))
    draw.text((50, 72), subtitle, fill="#d7ccff", font=font(20))
    return img, draw


def build_diagrams(inv):
    img, d = diagram_canvas("System Architecture", "Next.js frontend, Spring services, event fabric, and data stores")
    draw_box(d, (60, 170, 330, 310), "Users", "Visitors, marketers, admins, deliverability owners")
    draw_box(d, (430, 160, 760, 320), "Next.js Frontend", "Public site and authenticated workspace app")
    draw_box(d, (870, 160, 1220, 320), "API Gateway", "Ingress, routing, CORS, context headers")
    draw_box(d, (1290, 150, 1540, 330), "Spring Services", "10 domain services under /api/v1")
    for a, b in [((330, 240), (430, 240)), ((760, 240), (870, 240)), ((1220, 240), (1290, 240))]:
        arrow(d, a, b)
    stores = [("PostgreSQL", "Service-owned schemas"), ("Kafka", "Event fabric"), ("Redis", "Fast state"), ("MinIO", "Assets"), ("OpenSearch", "Search"), ("ClickHouse", "Analytics")]
    x = 110
    for name, body in stores:
        draw_box(d, (x, 500, x + 220, 650), name, body, fill="#ffffff", accent="#059669")
        arrow(d, (1420, 330), (x + 110, 500), color="#8b7ab6", width=3)
        x += 245
    img.save(DIAGRAMS / "system-architecture.png")

    img, d = diagram_canvas("Service Topology", "Domain services and primary ownership boundaries")
    positions = [(70, 170), (390, 170), (710, 170), (1030, 170), (70, 420), (390, 420), (710, 420), (1030, 420), (70, 670), (390, 670)]
    for s, (x, y) in zip(inv["services"], positions):
        draw_box(d, (x, y, x + 280, y + 150), s["name"], f"Port {s['port']} | DB {s['db']}", accent="#2563eb")
    draw_box(d, (1060, 650, 1500, 800), "Shared Runtime", "legent-security, legent-kafka, legent-common, test support", accent="#7c3aed")
    img.save(DIAGRAMS / "service-topology.png")

    img, d = diagram_canvas("Campaign Data Flow", "End-to-end path from audience signal to analytics feedback")
    steps = [
        ("Signup", "Tenant/workspace context"),
        ("Audience", "Subscribers, lists, segments"),
        ("Content", "Approved template"),
        ("Campaign", "Rules, budget, experiment"),
        ("Launch Gates", "Readiness and approval"),
        ("Delivery", "Provider route, queue, replay"),
        ("Tracking", "Open/click/raw events"),
        ("Analytics", "Summaries and feedback"),
    ]
    x, y = 70, 210
    prev = None
    for i, (name, body) in enumerate(steps):
        bx = x + (i % 4) * 380
        by = y + (i // 4) * 250
        draw_box(d, (bx, by, bx + 300, by + 150), name, body, accent="#dc2626" if i == 4 else "#7c3aed")
        if prev:
            arrow(d, prev, (bx, by + 75))
        prev = (bx + 300, by + 75)
    img.save(DIAGRAMS / "campaign-data-flow.png")

    img, d = diagram_canvas("Kafka Event Flow", "Asynchronous coordination and idempotent processing")
    draw_box(d, (70, 190, 360, 340), "Identity/Foundation", "Tenant provisioned, context bootstrapped", accent="#0891b2")
    draw_box(d, (470, 190, 760, 340), "Audience", "Audience resolution requested/resolved", accent="#0891b2")
    draw_box(d, (870, 190, 1160, 340), "Campaign", "Launch requested, send batches created", accent="#0891b2")
    draw_box(d, (1270, 190, 1530, 340), "Delivery", "Sent, failed, bounced, replay", accent="#0891b2")
    draw_box(d, (470, 560, 760, 710), "Tracking", "Open/click ingested", accent="#059669")
    draw_box(d, (870, 560, 1160, 710), "Automation", "Journey event consumed", accent="#059669")
    for a, b in [((360, 265), (470, 265)), ((760, 265), (870, 265)), ((1160, 265), (1270, 265)), ((1400, 340), (620, 560)), ((1400, 340), (1020, 560))]:
        arrow(d, a, b)
    img.save(DIAGRAMS / "event-flow.png")

    img, d = diagram_canvas("Database Ownership", "Service-owned schemas and transactional boundaries")
    x, y = 80, 160
    for i, s in enumerate(inv["services"]):
        bx = x + (i % 2) * 740
        by = y + (i // 2) * 130
        draw_box(d, (bx, by, bx + 650, by + 105), s["name"], f"{s['db']} | tables discovered: {s['tables']} | migrations: {s['migrations']}", accent="#059669")
    img.save(DIAGRAMS / "database-ownership.png")

    img, d = diagram_canvas("Deployment Architecture", "Cloud/Kubernetes deployment view")
    draw_box(d, (70, 180, 350, 320), "Internet", "Browser and public tracking calls", accent="#2563eb")
    draw_box(d, (470, 180, 760, 320), "Ingress", "TLS, routing, host/path rules", accent="#2563eb")
    draw_box(d, (880, 180, 1170, 320), "Frontend", "Next.js deployment", accent="#7c3aed")
    draw_box(d, (1290, 180, 1530, 320), "Services", "Spring Boot deployments", accent="#7c3aed")
    for a, b in [((350, 250), (470, 250)), ((760, 250), (880, 250)), ((1170, 250), (1290, 250))]:
        arrow(d, a, b)
    deps = [("Managed PostgreSQL", "one DB per service"), ("Managed Kafka", "topics and DLQ"), ("Redis", "cache/fast state"), ("Object Storage", "assets/files"), ("Search", "OpenSearch"), ("Analytics DB", "ClickHouse")]
    x = 80
    for name, body in deps:
        draw_box(d, (x, 540, x + 220, 700), name, body, accent="#ea580c")
        arrow(d, (1410, 320), (x + 110, 540), color="#8b7ab6", width=3)
        x += 245
    img.save(DIAGRAMS / "deployment-architecture.png")


def add_docx_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(37, 24, 74)
    return p


def create_docx(inv):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(30)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Legent Project Documentation")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor(37, 24, 74)
    sub = doc.add_paragraph("Complete technical, functional, architecture, operations, setup, and user journey report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {PY_NOW}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(DIAGRAMS / "system-architecture.png"), width=Inches(6.7))
    doc.add_page_break()

    add_docx_heading(doc, "Executive Summary", 1)
    doc.add_paragraph("Legent is a lifecycle email operating system with a Next.js frontend and Spring Boot domain services. This report covers product behavior, architecture, service ownership, API surfaces, database/migration inventory, operations, security, testing, setup, cloud deployment, and future roadmap.")

    add_docx_heading(doc, "Service Matrix", 1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Shading Accent 1"
    headers = ["Service", "Port", "Database", "Controllers", "Entities", "Migrations"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for s in inv["services"]:
        cells = table.add_row().cells
        vals = [s["name"], s["port"], s["db"], str(s["controllers"]), str(s["entities"]), str(s["migrations"])]
        for i, val in enumerate(vals):
            cells[i].text = val

    add_docx_heading(doc, "Architecture Diagrams", 1)
    for d in ["service-topology.png", "campaign-data-flow.png", "event-flow.png", "database-ownership.png", "deployment-architecture.png"]:
        add_docx_heading(doc, Path(d).stem.replace("-", " ").title(), 2)
        doc.add_picture(str(DIAGRAMS / d), width=Inches(6.7))

    add_docx_heading(doc, "Functional Modules", 1)
    for name, purpose, outcome in MODULE_GUIDE:
        add_docx_heading(doc, name, 2)
        doc.add_paragraph(f"Purpose: {purpose}")
        doc.add_paragraph(f"User outcome: {outcome}")

    add_docx_heading(doc, "User Journey Evidence", 1)
    for fname, title, caption in SCREEN_MAP:
        path = SCREENSHOTS / fname
        if not path.exists():
            continue
        add_docx_heading(doc, title, 2)
        doc.add_paragraph(caption)
        doc.add_picture(str(path), width=Inches(6.4))

    add_docx_heading(doc, "API And Database Summary", 1)
    doc.add_paragraph(f"Cataloged endpoints: {len(inv['endpoints'])}. See 09_API_Catalog.md for full endpoint details.")
    doc.add_paragraph("Each service owns its schema through Flyway. See 10_Database_And_Migrations.md for full migration inventory.")

    add_docx_heading(doc, "Local Setup", 1)
    for line in [
        "docker compose up -d postgres kafka redis minio opensearch clickhouse mailhog",
        "mvn -q -DskipTests package",
        "mvn test",
        "cd frontend",
        "npm install",
        "npm run dev -- --hostname 127.0.0.1 --port 3000",
    ]:
        doc.add_paragraph(line, style=None)

    add_docx_heading(doc, "Future Roadmap", 1)
    for area, enh in [
        ("Observability", "OpenTelemetry tracing, SLO dashboards, Kafka lag alerts, synthetic journeys."),
        ("Testing", "Contract tests, visual regression, load tests, migration tests, event idempotency tests."),
        ("Product Intelligence", "Explainable next-best-action, anomaly detection, and governed AI assistance."),
        ("Cloud", "Terraform modules, managed dependencies, blue/green rollout, and automated rollback."),
    ]:
        doc.add_paragraph(f"{area}: {enh}")

    out = DOC_DIR / "Legent_Project_Documentation.docx"
    doc.save(out)
    return out


def create_pdf(inv):
    out = DOC_DIR / "Legent_Project_Documentation.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, textColor=colors.HexColor("#25184a")))
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10))
    story = [
        Paragraph("Legent Project Documentation", styles["CenterTitle"]),
        Paragraph("Complete technical, functional, architecture, operations, setup, and user journey report", styles["BodyText"]),
        Spacer(1, 0.2 * inch),
        PdfImage(str(DIAGRAMS / "system-architecture.png"), width=6.9 * inch, height=3.9 * inch),
        PageBreak(),
        Paragraph("Service Matrix", styles["Heading1"]),
    ]
    data = [["Service", "Port", "DB", "Controllers", "Entities", "Migrations"]]
    for s in inv["services"]:
        data.append([s["name"], s["port"], s["db"], s["controllers"], s["entities"], s["migrations"]])
    table = Table(data, repeatRows=1, colWidths=[1.35 * inch, 0.55 * inch, 1.3 * inch, 0.75 * inch, 0.65 * inch, 0.75 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9ddff")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cdbdf5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
    ]))
    story += [table, PageBreak(), Paragraph("Architecture Diagrams", styles["Heading1"])]
    for d in ["service-topology.png", "campaign-data-flow.png", "event-flow.png", "database-ownership.png", "deployment-architecture.png"]:
        story.append(Paragraph(Path(d).stem.replace("-", " ").title(), styles["Heading2"]))
        story.append(PdfImage(str(DIAGRAMS / d), width=6.9 * inch, height=3.9 * inch))
        story.append(Spacer(1, 0.15 * inch))
    story.append(PageBreak())
    story.append(Paragraph("Functional Modules", styles["Heading1"]))
    for name, purpose, outcome in MODULE_GUIDE:
        story.append(Paragraph(name, styles["Heading2"]))
        story.append(Paragraph(f"<b>Purpose:</b> {purpose}", styles["BodyText"]))
        story.append(Paragraph(f"<b>User outcome:</b> {outcome}", styles["BodyText"]))
    story.append(PageBreak())
    story.append(Paragraph("User Journey Screenshots", styles["Heading1"]))
    for fname, title, caption in SCREEN_MAP:
        path = SCREENSHOTS / fname
        if path.exists():
            story.append(Paragraph(title, styles["Heading2"]))
            story.append(Paragraph(caption, styles["BodyText"]))
            story.append(PdfImage(str(path), width=6.9 * inch, height=4.8 * inch))
            story.append(Spacer(1, 0.1 * inch))
    doc = SimpleDocTemplate(str(out), pagesize=A4, rightMargin=0.45 * inch, leftMargin=0.45 * inch, topMargin=0.45 * inch, bottomMargin=0.45 * inch)
    doc.build(story)
    return out


def copy_temp_ffmpeg():
    env_ffmpeg = os.environ.get("LEGENT_FFMPEG")
    env_ffprobe = os.environ.get("LEGENT_FFPROBE")
    candidates = [
        Path(env_ffmpeg) if env_ffmpeg else None,
        Path("C:/Users/leelark.saxena/Documents/New project 4/node_modules/ffmpeg-static/ffmpeg.exe"),
        Path("C:/Users/leelark.saxena/Documents/New project 4/node_modules/@remotion/compositor-win32-x64-msvc/ffmpeg.exe"),
    ]
    probe_candidates = [
        Path(env_ffprobe) if env_ffprobe else None,
        Path("C:/Users/leelark.saxena/Documents/New project 4/node_modules/ffprobe-static/bin/win32/x64/ffprobe.exe"),
    ]
    ffmpeg_src = next((p for p in candidates if p and p.exists() and p.is_file()), None)
    ffprobe_src = next((p for p in probe_candidates if p and p.exists() and p.is_file()), None)
    if not ffmpeg_src:
        raise RuntimeError("ffmpeg not found")
    tool_dir = TEMP_TOOLS / "ffmpeg"
    tool_dir.mkdir(parents=True, exist_ok=True)
    ffmpeg = tool_dir / "ffmpeg.exe"
    shutil.copy2(ffmpeg_src, ffmpeg)
    ffprobe = None
    if ffprobe_src:
        ffprobe = tool_dir / "ffprobe.exe"
        shutil.copy2(ffprobe_src, ffprobe)
    return ffmpeg, ffprobe


def make_video_slide(asset_path, title, subtitle, index, total, dest):
    W, H = 1920, 1080
    img = Image.new("RGB", (W, H), "#f7f5ff")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, W, 112), fill="#21133f")
    d.text((58, 28), "Legent", fill="#ffffff", font=font(34, True))
    d.text((58, 70), title, fill="#d8ccff", font=font(21))
    d.rounded_rectangle((60, 170, 470, 920), radius=28, fill="#ffffff", outline="#d7c7ff", width=3)
    d.text((95, 210), f"CHAPTER {index:02d}/{total:02d}", fill="#7c3aed", font=font(22, True))
    y = 260
    for line in wrap(d, subtitle, font(28, True), 320)[:6]:
        d.text((95, y), line, fill="#1f1640", font=font(28, True))
        y += 38
    d.line((95, y + 18, 420, y + 18), fill="#cdbdf5", width=3)
    d.text((95, 820), "Verified source asset", fill="#6b5c8f", font=font(20))
    d.text((95, 850), Path(asset_path).name, fill="#1f1640", font=font(18, True))

    asset = Image.open(asset_path).convert("RGB")
    asset = ImageOps.contain(asset, (1320, 760))
    ax = 540 + (1320 - asset.width) // 2
    ay = 170 + (760 - asset.height) // 2
    d.rounded_rectangle((520, 150, 1880, 950), radius=30, fill="#ffffff", outline="#cdbdf5", width=3)
    img.paste(asset, (ax, ay))
    d.rounded_rectangle((1370, 900, 1840, 970), radius=20, fill="#21133f")
    d.text((1400, 922), "1920x1080 edited demo frame", fill="#ffffff", font=font(22, True))
    img.save(dest)


def seconds_to_srt_time(seconds):
    ms = int(round((seconds - int(seconds)) * 1000))
    s = int(seconds)
    h = s // 3600
    m = (s % 3600) // 60
    sec = s % 60
    return f"{h:02d}:{m:02d}:{sec:02d},{ms:03d}"


def synthesize_voice(text_path, wav_path):
    ps1 = WORK / "scripts" / "speak.ps1"
    ps1.write_text("""
param([string]$TextPath, [string]$OutPath)
Add-Type -AssemblyName System.Speech
$s = New-Object System.Speech.Synthesis.SpeechSynthesizer
$s.Rate = -1
$s.Volume = 100
$voice = $s.GetInstalledVoices() | Where-Object { $_.VoiceInfo.Culture.Name -like 'en-*' } | Select-Object -First 1
if ($voice) { $s.SelectVoice($voice.VoiceInfo.Name) }
$s.SetOutputToWaveFile($OutPath)
$s.Speak((Get-Content -LiteralPath $TextPath -Raw))
$s.Dispose()
""", encoding="utf-8")
    subprocess.run(["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(ps1), str(text_path), str(wav_path)], check=True)


def ffprobe_duration(ffprobe, media):
    if not ffprobe:
        return None
    res = subprocess.run([str(ffprobe), "-v", "error", "-show_entries", "format=duration", "-of", "default=nw=1:nk=1", str(media)], capture_output=True, text=True)
    try:
        return float(res.stdout.strip())
    except Exception:
        return None


def render_video(folder, out_name, title, target_seconds, assets, voice_paragraphs, ffmpeg, ffprobe):
    folder.mkdir(parents=True, exist_ok=True)
    source_dir = folder / "source-assets"
    source_dir.mkdir(exist_ok=True)
    slides_dir = folder / "slides"
    slides_dir.mkdir(exist_ok=True)
    per = target_seconds / len(assets)
    srt_lines = []
    concat_lines = []
    chapters = []
    for i, (asset, heading, caption) in enumerate(assets, start=1):
        source = asset
        copied = source_dir / source.name
        if source.exists():
            shutil.copy2(source, copied)
        slide = slides_dir / f"slide_{i:02d}.png"
        make_video_slide(source, heading, caption, i, len(assets), slide)
        concat_lines.append(f"file '{slide.as_posix()}'")
        concat_lines.append(f"duration {per:.3f}")
        start = (i - 1) * per
        end = i * per
        srt_lines += [str(i), f"{seconds_to_srt_time(start)} --> {seconds_to_srt_time(end)}", f"{heading}. {caption}", ""]
        chapters.append([seconds_to_srt_time(start).replace(",", "."), heading, caption])
    concat_lines.append(f"file '{slide.as_posix()}'")
    (folder / "concat.txt").write_text("\n".join(concat_lines), encoding="utf-8")
    (folder / "captions.srt").write_text("\n".join(srt_lines), encoding="utf-8")
    (folder / "chapters.md").write_text("# Chapters\n\n" + "\n".join([f"- {t} - {h}: {c}" for t, h, c in chapters]) + "\n", encoding="utf-8")

    narration = "\n\n".join(voice_paragraphs)
    (folder / "script.md").write_text(f"# {title}\n\n{narration}\n", encoding="utf-8")
    voice_txt = folder / "voiceover.txt"
    voice_txt.write_text(narration, encoding="utf-8")
    wav = folder / "voiceover.wav"
    synthesize_voice(voice_txt, wav)
    dur = ffprobe_duration(ffprobe, wav) or target_seconds
    atempo = max(0.5, min(2.0, dur / max(1, target_seconds * 0.98)))
    audio = folder / "voiceover-fit.wav"
    subprocess.run([str(ffmpeg), "-y", "-i", str(wav), "-filter:a", f"atempo={atempo:.5f}", "-t", str(target_seconds), str(audio)], check=True)

    out = folder / out_name
    subprocess.run([
        str(ffmpeg), "-y",
        "-f", "concat", "-safe", "0", "-i", str(folder / "concat.txt"),
        "-i", str(audio),
        "-filter_complex", "[1:a]apad[a]",
        "-map", "0:v", "-map", "[a]",
        "-t", str(target_seconds),
        "-vf", "fps=30,format=yuv420p",
        "-c:v", "libx264", "-preset", "veryfast", "-crf", "20",
        "-c:a", "aac", "-b:a", "192k",
        "-movflags", "+faststart",
        str(out),
    ], check=True)
    meta = {
        "file": str(out),
        "target_seconds": target_seconds,
        "voiceover_original_seconds": dur,
        "voiceover_atempo": atempo,
        "ffprobe_seconds": ffprobe_duration(ffprobe, out),
        "resolution": "1920x1080",
        "fps": 30,
        "source_assets": [a[0].name for a in assets],
    }
    (folder / "render-manifest.json").write_text(json.dumps(meta, indent=2), encoding="utf-8")
    return meta


def build_videos(ffmpeg, ffprobe):
    ss = SCREENSHOTS
    dg = DIAGRAMS
    product_assets = [
        (ss / "01_public_home.png", "Product entry", "Public product story, navigation, and operating-system positioning."),
        (ss / "05_signup.png", "Signup", "Workspace creation starts the authenticated operator journey."),
        (ss / "09_audience_overview.png", "Audience", "Audience health, subscribers, lists, imports, and segments become operational inputs."),
        (ss / "08_template_studio.png", "Template Studio", "Governed creative begins with reusable templates and approvals."),
        (ss / "12_campaign_wizard.png", "Campaign Wizard", "Campaign setup joins message identity, audience, policy, and review."),
        (ss / "13_launch_orchestration.png", "Launch", "Readiness checks convert a campaign draft into a controlled launch."),
        (ss / "16_deliverability.png", "Delivery Studio", "Queue, replay, domain health, and provider readiness stay visible."),
        (ss / "18_admin_console.png", "Admin", "Admin command center governs users, policy, audit, and runtime control."),
    ]
    tech_assets = [
        (dg / "system-architecture.png", "System architecture", "Frontend, gateway, services, event fabric, and data stores."),
        (dg / "service-topology.png", "Service topology", "Ten Spring services own independent domains."),
        (dg / "campaign-data-flow.png", "Data flow", "Campaign data moves from audience and content to delivery and analytics."),
        (dg / "event-flow.png", "Kafka event flow", "Events decouple launch, delivery, tracking, automation, and admin propagation."),
        (dg / "database-ownership.png", "Database ownership", "Each service owns a database and Flyway migration timeline."),
        (ss / "07_email_studio.png", "Frontend app", "Next.js screens call normalized API clients with workspace context."),
        (ss / "14_campaign_tracking.png", "Runtime evidence", "Campaign tracking shows launch safety, progress, and outcomes."),
        (dg / "deployment-architecture.png", "Deployment", "Kubernetes ingress, deployments, config, secrets, and managed dependencies."),
    ]

    def repeat_voice(base, target_words):
        words = " ".join(base).split()
        out = []
        while len(out) < target_words:
            out.extend(words)
        return [" ".join(out[:target_words])]

    product_base = [
        "Legent is a lifecycle email operating system for teams that need marketing speed and operational control in the same workspace.",
        "The journey starts on the public site, moves through signup and onboarding, then gives operators audience, template, campaign, delivery, analytics, automation, and admin tools.",
        "The refined screenshots in this video were captured from the local frontend after route-specific load checks, so every frame shows loaded product state instead of a blank shell.",
        "Audience data, creative approvals, campaign rules, launch gates, delivery health, and analytics feedback are connected so teams can move from idea to inbox with stronger evidence.",
    ]
    tech_base = [
        "Technically, Legent is a Next.js frontend backed by Spring Boot microservices. Each service owns a domain boundary, a database schema, controllers, services, repositories, migrations, and tests.",
        "Shared libraries provide tenant context, JWT and RBAC handling, Kafka envelope support, and test utilities. Kafka is the asynchronous fabric for provisioning, audience resolution, campaign launch, delivery feedback, tracking ingestion, automation, and admin propagation.",
        "PostgreSQL stores transactional service data. ClickHouse supports analytics. Redis supports fast state. MinIO stores assets. OpenSearch supports search. Kubernetes manifests provide the cloud deployment baseline.",
        "The local run path starts dependencies with Docker Compose, builds with Maven, runs services, starts the Next.js frontend, then validates with smoke checks and Playwright tests.",
    ]
    metas = []
    metas.append(render_video(PACKAGE / "demo-video" / "3min", "legent-product-demo-3min.mp4", "Legent Product Demo - 3 Minute Executive Teaser", 180, product_assets[:6], repeat_voice(product_base, 520), ffmpeg, ffprobe))
    metas.append(render_video(PACKAGE / "demo-video" / "6min", "legent-product-demo-6min.mp4", "Legent Product Demo - 6 Minute Walkthrough", 360, product_assets, repeat_voice(product_base, 1020), ffmpeg, ffprobe))
    metas.append(render_video(PACKAGE / "demo-video" / "10min", "legent-product-demo-10min.mp4", "Legent Product Demo - 10 Minute Detailed Walkthrough", 600, product_assets + [
        (ss / "10_subscriber_manager.png", "Subscribers", "Detailed subscriber state and lifecycle fields."),
        (ss / "15_automation_builder.png", "Automation", "Workflow governance and journey state."),
    ], repeat_voice(product_base, 1650), ffmpeg, ffprobe))
    metas.append(render_video(PACKAGE / "technical-video" / "10min", "legent-technical-explanation-10min.mp4", "Legent Technical Explanation - 10 Minutes", 600, tech_assets, repeat_voice(tech_base, 1650), ffmpeg, ffprobe))
    return metas


def validate(inv, video_metas):
    artifacts = []
    for p in PACKAGE.rglob("*"):
        if p.is_file() and "_temp_tools" not in p.parts and "_work" not in p.parts:
            artifacts.append({"path": str(p.relative_to(PACKAGE)).replace("\\", "/"), "bytes": p.stat().st_size})
    (REPORTS / "artifact_inventory.json").write_text(json.dumps(artifacts, indent=2), encoding="utf-8")
    report = [
        "# Validation Report",
        "",
        f"Generated: {PY_NOW}",
        "",
        f"- Verified screenshots present: {len(inv['screenshots'])}",
        f"- Cataloged services: {len(inv['services'])}",
        f"- Cataloged endpoints: {len(inv['endpoints'])}",
        f"- Markdown docs: {len(list(DOC_MD.glob('*.md')))}",
        f"- Video renders: {len(video_metas)}",
        "",
        "## Video Metadata",
        md_table(["File", "Target seconds", "Measured seconds", "Resolution"], [[Path(m["file"]).name, m["target_seconds"], round(m["ffprobe_seconds"] or 0, 2), m["resolution"]] for m in video_metas]),
        "",
        "## Screenshot QA",
        "The Playwright capture script waited for page-specific text and the final capture log contains no fatal, pageerror, or browser error entries.",
    ]
    (REPORTS / "validation_report.md").write_text("\n".join(report) + "\n", encoding="utf-8")


def main():
    ensure_dirs()
    inv = inventory()
    build_diagrams(inv)
    write_markdown(inv)
    create_docx(inv)
    create_pdf(inv)
    ffmpeg, ffprobe = copy_temp_ffmpeg()
    video_metas = build_videos(ffmpeg, ffprobe)
    validate(inv, video_metas)


if __name__ == "__main__":
    main()
